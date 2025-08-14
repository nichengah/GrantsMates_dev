"""
Author: Yuewen Li
Date: 2025/07/25
Description: this script adds the actual URL in brackets after any existing hyperlinks in a Word (.docx) document.

Usage: Set the path of the input file like this:
    input_file = r"D:/下载/UAlbany_Research_Contact_Detailed.docx".
    
Requirements:
Install the `python-docx` library using the command: 
    pip install python-docx
    

"""



from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.oxml import OxmlElement
import os
import sys
import re


def add_url_parentheses_directly_after_link(doc_path, output_path):
    """
    将URL括号直接添加在超链接文本后面 - 增强版
    """
    doc = Document(doc_path)
    total_hyperlinks = 0
    processed_hyperlinks = 0
    skipped_hyperlinks = 0

    print(f"\n{'=' * 50}")
    print(f"开始处理文档: {doc_path}")

    def insert_url_after_hyperlink(paragraph, hyperlink, url, location):
        """在超链接后直接插入URL - 增强版"""
        # 尝试找到超链接内的文本内容
        link_text = ""
        for elem in hyperlink.iterchildren():
            if elem.tag.endswith('r'):  # 文本run
                for t in elem.iterchildren():
                    if t.tag.endswith('t'):  # 文本元素
                        if t.text:
                            link_text += t.text
                    elif t.tag.endswith('tab'):  # 制表符
                        link_text += '\t'
                    elif t.tag.endswith('br'):  # 换行符
                        link_text += '\n'

        # 如果没有找到文本内容，尝试其他方法
        if not link_text:
            # 尝试从超链接属性中获取文本
            instr_text = hyperlink.get(qn('w:instr')) or ""
            if 'HYPERLINK' in instr_text:
                # 尝试从域代码中提取文本
                match = re.search(r'\\o "(.*?)"', instr_text)
                if match:
                    link_text = match.group(1)
                else:
                    # 作为最后手段，使用URL作为文本
                    link_text = url
            else:
                link_text = "超链接"

        print(f"🔍 [{location}] 超链接文本: '{link_text}'")

        # 创建新run并插入
        new_run = OxmlElement('w:r')

        # 添加文本属性
        rPr = OxmlElement('w:rPr')
        new_run.append(rPr)

        # 添加文本内容
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')  # 保留空格
        t.text = f" ({url})"
        new_run.append(t)

        # 设置样式
        if rPr is not None:
            # 设置字体大小
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), "18")  # 9pt * 2 = 18
            rPr.append(sz)

            # 移除下划线
            u = OxmlElement('w:u')
            u.set(qn('w:val'), "none")
            rPr.append(u)

        # 在超链接元素后插入新run
        hyperlink.addnext(new_run)
        return True

    def process_hyperlink(paragraph, hyperlink, location):
        nonlocal total_hyperlinks, processed_hyperlinks, skipped_hyperlinks
        total_hyperlinks += 1

        # 获取关系ID
        rel_id = hyperlink.get(qn('r:id'))
        if not rel_id:
            print(f"⚠️ [{location}] 跳过无rel_id的超链接")
            skipped_hyperlinks += 1
            return

        # 获取真实URL
        try:
            if rel_id in doc.part.rels:
                url = doc.part.rels[rel_id]._target
            else:
                # 尝试从超链接属性中获取URL
                if hyperlink.get(qn('w:instr')):
                    instr_text = hyperlink.get(qn('w:instr'))
                    if 'HYPERLINK' in instr_text:
                        match = re.search(r'"([^"]+)"', instr_text)
                        if match:
                            url = match.group(1)
                        else:
                            print(f"❌ [{location}] 无法从域代码中提取URL")
                            skipped_hyperlinks += 1
                            return
                else:
                    print(f"❌ [{location}] rel_id={rel_id} 未找到对应关系")
                    skipped_hyperlinks += 1
                    return
        except Exception as e:
            print(f"❌ [{location}] 获取URL时出错: {str(e)}")
            skipped_hyperlinks += 1
            return

        # 在超链接后直接插入URL
        if insert_url_after_hyperlink(paragraph, hyperlink, url, location):
            print(f"✅ [{location}] 已处理: 添加 ({url})")
            processed_hyperlinks += 1
        else:
            skipped_hyperlinks += 1

    # 处理正文段落
    for i, paragraph in enumerate(doc.paragraphs):
        # 使用更可靠的XPath查询
        hyperlinks = paragraph._element.xpath('.//w:hyperlink|.//w:fldSimple[contains(@w:instr, "HYPERLINK")]')
        for j, hyperlink in enumerate(hyperlinks):
            location = f"段落 {i + 1}.{j + 1}"
            process_hyperlink(paragraph, hyperlink, location)

    # 处理表格
    for t, table in enumerate(doc.tables):
        for r, row in enumerate(table.rows):
            for c, cell in enumerate(row.cells):
                for p, paragraph in enumerate(cell.paragraphs):
                    hyperlinks = paragraph._element.xpath(
                        './/w:hyperlink|.//w:fldSimple[contains(@w:instr, "HYPERLINK")]')
                    for h, hyperlink in enumerate(hyperlinks):
                        location = f"表格 {t + 1}行{r + 1}列{c + 1}段落{p + 1}.{h + 1}"
                        process_hyperlink(paragraph, hyperlink, location)

    # 处理页眉
    for s, section in enumerate(doc.sections):
        if section.header:
            for i, paragraph in enumerate(section.header.paragraphs):
                hyperlinks = paragraph._element.xpath('.//w:hyperlink|.//w:fldSimple[contains(@w:instr, "HYPERLINK")]')
                for j, hyperlink in enumerate(hyperlinks):
                    location = f"页眉 {s + 1}段落{i + 1}.{j + 1}"
                    process_hyperlink(paragraph, hyperlink, location)

    # 处理页脚
    for s, section in enumerate(doc.sections):
        if section.footer:
            for i, paragraph in enumerate(section.footer.paragraphs):
                hyperlinks = paragraph._element.xpath('.//w:hyperlink|.//w:fldSimple[contains(@w:instr, "HYPERLINK")]')
                for j, hyperlink in enumerate(hyperlinks):
                    location = f"页脚 {s + 1}段落{i + 1}.{j + 1}"
                    process_hyperlink(paragraph, hyperlink, location)

    # 保存文档
    doc.save(output_path)

    print(f"\n处理摘要:")
    print(f"总超链接数: {total_hyperlinks}")
    print(f"已处理: {processed_hyperlinks}")
    print(f"已跳过: {skipped_hyperlinks}")
    print(f"输出文件: {output_path}")
    print(f"{'=' * 50}\n")

    return total_hyperlinks > 0


def create_direct_placement_output(input_path, suffix="_direct_urls"):
    """创建直接定位URL的输出文件"""
    if not os.path.exists(input_path):
        print(f"❌ 文件不存在: {input_path}")
        return None

    # 生成输出路径
    base, ext = os.path.splitext(input_path)
    output_path = f"{base}{suffix}{ext}"

    # 避免文件名冲突
    counter = 1
    while os.path.exists(output_path):
        output_path = f"{base}{suffix}_{counter}{ext}"
        counter += 1

    # 处理文档
    found_hyperlinks = add_url_parentheses_directly_after_link(input_path, output_path)

    if not found_hyperlinks:
        print("\n⚠️ 警告: 未在文档中找到任何超链接")

    return output_path


if __name__ == "__main__":
    # 获取当前脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    # script_dir = os.getcwd()

    # 输入文件路径
    input_file = r"C:/Users/李月雯/PycharmProjects/PythonProject1/GrantsMates_dev/add_hyperlinks/AWARD_REVIEW_&_NEGOTIATION.docx"

    print(f"当前工作目录: {os.getcwd()}")
    print(f"脚本目录: {script_dir}")
    print(f"输入文件: {input_file}")

    # 创建新文件并处理
    output_file = create_direct_placement_output(input_file)

    if output_file:
        print(f"✅ 处理完成! 输出文件: {output_file}")
        print("URL括号已直接添加在超链接文本后面")

        # 在Windows上自动打开文件
        if sys.platform == 'win32':
            os.startfile(output_file)
    else:
        print("❌ 处理失败")
